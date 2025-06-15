from flask import Flask, request, render_template, send_file, jsonify
import os
import io
import base64
import re
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import google.generativeai as genai
from werkzeug.utils import secure_filename
import tempfile
import zipfile
from datetime import datetime

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Gemini API設定
genai.configure(api_key=os.environ.get('GEMINI_API_KEY'))
model = genai.GenerativeModel('gemini-1.5-flash')

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'png', 'jpg', 'jpeg', 'gif', 'bmp'}

def extract_text_from_image(image_path):
    """Gemini APIを使用して画像からテキストを抽出"""
    try:
        with open(image_path, 'rb') as image_file:
            image_data = image_file.read()
        
        # Gemini APIでOCR
        image = Image.open(io.BytesIO(image_data))
        
        prompt = """
        この画像から英語のテキストを正確に抽出してください。
        レイアウトや改行を可能な限り保持し、読みやすい形で出力してください。
        テキストのみを返してください。
        """
        
        response = model.generate_content([prompt, image])
        return response.text.strip()
    
    except Exception as e:
        print(f"OCRエラー: {e}")
        return ""

def translate_text(text):
    """テキストを日本語に翻訳"""
    try:
        prompt = f"""
        以下の英語テキストを自然で読みやすい日本語に翻訳してください。
        文学的な表現や専門用語も適切に翻訳し、原文の意味とニュアンスを保持してください。

        英語テキスト:
        {text}
        """
        
        response = model.generate_content(prompt)
        return response.text.strip()
    
    except Exception as e:
        print(f"翻訳エラー: {e}")
        return text

def extract_important_words(text):
    """重要な単語を抽出し、例文とともに解説を生成"""
    try:
        prompt = f"""
        以下の英語テキストから、学習に重要な単語を最大20個選択し、
        各単語について以下の形式でJSONで返してください：

        {{
            "words": [
                {{
                    "word": "単語",
                    "definition": "日本語での意味・定義",
                    "example": "その単語を使った例文（英語）",
                    "example_translation": "例文の日本語訳",
                    "level": "初級/中級/上級"
                }}
            ]
        }}

        選択基準：
        - 頻出度が高い重要語彙
        - 学術的・専門的な語彙
        - 文脈上重要な意味を持つ語彙
        - 学習者にとって有益な語彙

        英語テキスト:
        {text[:2000]}  # テキストが長すぎる場合は最初の2000文字のみ
        """
        
        response = model.generate_content(prompt)
        
        # JSONレスポンスから辞書を抽出
        import json
        try:
            # ```json ``` で囲まれている場合の処理
            response_text = response.text.strip()
            if "```json" in response_text:
                json_start = response_text.find("```json") + 7
                json_end = response_text.find("```", json_start)
                json_text = response_text[json_start:json_end].strip()
            else:
                json_text = response_text
            
            data = json.loads(json_text)
            return data.get("words", [])
        except json.JSONDecodeError:
            print("JSON解析エラー")
            return []
    
    except Exception as e:
        print(f"単語抽出エラー: {e}")
        return []

def analyze_font_style(image_path):
    """画像からフォントスタイルを分析"""
    try:
        prompt = """
        この画像のテキストのフォントスタイルを分析し、
        以下のカテゴリから最も近いものを選択してください：
        - serif (明朝体系)
        - sans-serif (ゴシック体系)
        - monospace (等幅)
        - decorative (装飾的)
        
        また、以下の特徴も判定してください：
        - 太さ: thin/normal/bold
        - スタイル: normal/italic
        
        結果を「フォント種類,太さ,スタイル」の形式で返してください。
        例: serif,normal,normal
        """
        
        image = Image.open(image_path)
        response = model.generate_content([prompt, image])
        return response.text.strip()
    
    except Exception as e:
        print(f"フォント分析エラー: {e}")
        return "serif,normal,normal"

def create_word_document(original_text, translated_text, important_words, font_style):
    """Wordドキュメントを作成"""
    doc = Document()
    
    # フォントスタイルの設定
    font_info = font_style.split(',')
    font_family = "Times New Roman" if font_info[0] == "serif" else "Arial"
    is_bold = font_info[1] == "bold" if len(font_info) > 1 else False
    is_italic = font_info[2] == "italic" if len(font_info) > 2 else False
    
    # タイトル
    title = doc.add_heading('英語テキスト翻訳・単語解説', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 作成日時
    date_p = doc.add_paragraph(f'作成日時: {datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 原文セクション
    doc.add_heading('原文（英語）', level=1)
    original_p = doc.add_paragraph(original_text)
    for run in original_p.runs:
        run.font.name = font_family
        run.font.size = Pt(11)
        run.font.bold = is_bold
        run.font.italic = is_italic
    
    # 翻訳セクション
    doc.add_heading('翻訳（日本語）', level=1)
    translation_p = doc.add_paragraph(translated_text)
    for run in translation_p.runs:
        run.font.name = 'Yu Gothic'
        run.font.size = Pt(11)
    
    # 重要単語セクション
    doc.add_heading('重要単語解説', level=1)
    
    for i, word_info in enumerate(important_words, 1):
        # 単語見出し
        word_heading = doc.add_heading(f'{i}. {word_info.get("word", "")}', level=2)
        
        # 定義
        def_p = doc.add_paragraph()
        def_p.add_run('意味: ').bold = True
        def_p.add_run(word_info.get("definition", ""))
        
        # レベル
        level_p = doc.add_paragraph()
        level_p.add_run('レベル: ').bold = True
        level_p.add_run(word_info.get("level", ""))
        
        # 例文
        example_p = doc.add_paragraph()
        example_p.add_run('例文: ').bold = True
        example_p.add_run(word_info.get("example", ""))
        
        # 例文翻訳
        example_trans_p = doc.add_paragraph()
        example_trans_p.add_run('例文翻訳: ').bold = True
        example_trans_p.add_run(word_info.get("example_translation", ""))
        
        # 間隔調整
        doc.add_paragraph()
    
    return doc

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'ファイルが選択されていません'}), 400
    
    files = request.files.getlist('files')
    
    if len(files) > 20:
        return jsonify({'error': '最大20枚まで処理可能です'}), 400
    
    # 一時ディレクトリを作成
    temp_dir = tempfile.mkdtemp()
    uploaded_files = []
    
    try:
        # ファイルをアップロード
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(temp_dir, filename)
                file.save(filepath)
                uploaded_files.append(filepath)
        
        if not uploaded_files:
            return jsonify({'error': '有効な画像ファイルがありません'}), 400
        
        # OCR処理
        all_text = ""
        font_style = "serif,normal,normal"
        
        for image_path in uploaded_files:
            extracted_text = extract_text_from_image(image_path)
            if extracted_text:
                all_text += extracted_text + "\n\n"
            
            # 最初の画像からフォントスタイルを分析
            if image_path == uploaded_files[0]:
                font_style = analyze_font_style(image_path)
        
        if not all_text.strip():
            return jsonify({'error': 'テキストを抽出できませんでした'}), 400
        
        # 翻訳
        translated_text = translate_text(all_text)
        
        # 重要単語抽出
        important_words = extract_important_words(all_text)
        
        # Wordドキュメント作成
        doc = create_word_document(all_text, translated_text, important_words, font_style)
        
        # ファイル保存
        output_filename = f"translation_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(temp_dir, output_filename)
        doc.save(output_path)
        
        # ファイルをバイナリで読み込み
        with open(output_path, 'rb') as f:
            file_data = f.read()
        
        # 一時ファイル削除
        import shutil
        shutil.rmtree(temp_dir)
        
        # レスポンス
        return jsonify({
            'status': 'success',
            'original_text': all_text[:500] + '...' if len(all_text) > 500 else all_text,
            'translated_text': translated_text[:500] + '...' if len(translated_text) > 500 else translated_text,
            'word_count': len(important_words),
            'download_url': f'/download/{output_filename}',
            'file_data': base64.b64encode(file_data).decode('utf-8'),
            'filename': output_filename
        })
    
    except Exception as e:
        # エラー時は一時ディレクトリを削除
        import shutil
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        return jsonify({'error': f'処理中にエラーが発生しました: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    # この実装は簡略化されています
    # 実際にはセッションやデータベースを使用してファイルを管理する必要があります
    return jsonify({'error': 'ダイレクトダウンロードは実装されていません'}), 404

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)